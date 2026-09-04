import io
from typing import Optional


async def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except ImportError:
        return _fallback_pdf_extraction(file_bytes)
    except Exception as e:
        return f"Error extracting PDF text: {str(e)}"


async def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"Error extracting DOCX text: {str(e)}"


def _fallback_pdf_extraction(file_bytes: bytes) -> str:
    try:
        import pdf2image
        from PIL import Image
        return "PDF text extraction requires PyMuPDF. Install with: pip install PyMuPDF"
    except Exception:
        return "Unable to extract text from PDF. Please ensure PyMuPDF is installed."


def validate_resume_file(filename: str, file_size: int, max_size: int = 10 * 1024 * 1024) -> Optional[str]:
    if file_size > max_size:
        return f"File size exceeds {max_size // (1024*1024)}MB limit"
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext not in ('pdf', 'docx', 'doc'):
        return "Only PDF and DOCX files are supported"
    return None
